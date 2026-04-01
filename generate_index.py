#!/usr/bin/env python3
"""
docs-index-generator: 扫描 docs/ 目录，生成 index.html 索引页和 pages/ 目录下的独立 Document_Page HTML 文件。
"""

import os
import re
import html as html_module
import base64
import csv
import json
import mimetypes
from dataclasses import dataclass, field
from typing import List, Dict


# ============================================================
# 数据模型
# ============================================================

@dataclass
class FileNode:
    """文件或目录节点"""
    name: str               # 文件或目录名称
    path: str               # 相对于 docs/ 的路径
    is_dir: bool            # 是否为目录
    extension: str          # 文件扩展名（小写），目录为空字符串
    children: List['FileNode'] = field(default_factory=list)  # 子节点列表
    display_name: str = ''  # 显示名称（去掉扩展名）
    page_path: str = ''     # 对应的 Document_Page 相对路径，目录为空字符串


@dataclass
class FileTree:
    """文件树"""
    root: FileNode          # 根节点（docs/ 目录本身）
    file_count: int = 0     # 文件总数
    dir_count: int = 0      # 目录总数


@dataclass
class RenderResult:
    """渲染结果"""
    file_path: str          # 文件相对路径（相对于 docs/）
    html_content: str       # 渲染后的 HTML 片段（body 部分）
    file_type: str          # 文件类型标识
    size_bytes: int = 0     # 原始文件大小
    page_path: str = ''     # 输出的 Document_Page 相对路径（相对于 pages/）


@dataclass
class BuildStats:
    """构建统计"""
    total_files: int = 0            # 处理的文件总数
    total_dirs: int = 0             # 扫描的目录总数
    total_pages: int = 0            # 生成的 Document_Page 数量
    index_size_bytes: int = 0       # 生成的 index.html 大小
    pages_size_bytes: int = 0       # pages/ 目录总大小
    file_type_counts: Dict[str, int] = field(default_factory=dict)  # 各文件类型的数量统计
    errors: List[str] = field(default_factory=list)  # 处理过程中的错误列表


# ============================================================
# 文件类型到渲染器的映射
# ============================================================

RENDERER_MAP: Dict[str, str] = {
    # 文档格式
    '.md': '_render_markdown',
    '.pdf': '_render_pdf',
    '.html': '_render_html',
    '.htm': '_render_html',
    # 图片格式
    '.jpg': '_render_image',
    '.jpeg': '_render_image',
    '.png': '_render_image',
    '.gif': '_render_image',
    '.svg': '_render_image',
    '.bmp': '_render_image',
    '.webp': '_render_image',
    # Office 格式
    '.docx': '_render_docx',
    '.xlsx': '_render_xlsx',
    # 数据格式
    '.csv': '_render_csv',
    '.json': '_render_json',
    # 特殊格式
    '.mermaid': '_render_mermaid',
    '.tex': '_render_latex',
    '.latex': '_render_latex',
    # 代码文件 —— 使用 Pygments 进行语法高亮渲染（含行号）
    '.py': '_render_code',
    '.java': '_render_code',
    '.cpp': '_render_code',
    '.hpp': '_render_code',
    '.c': '_render_code',
    '.h': '_render_code',
    '.js': '_render_code',
    '.ts': '_render_code',
    '.go': '_render_code',
    '.rs': '_render_code',
    '.rb': '_render_code',
    '.php': '_render_code',
    '.swift': '_render_code',
    '.kt': '_render_code',
    '.scala': '_render_code',
    '.sh': '_render_code',
    '.bash': '_render_code',
    '.css': '_render_code',
    '.sql': '_render_code',
}
# 不在映射中的扩展名 -> 尝试纯文本渲染，失败则显示不支持提示


# ============================================================
# 核心类
# ============================================================

class DirectoryScanner:
    """递归扫描目录，构建文件树结构。"""

    def scan(self, root_path: str) -> FileTree:
        """递归扫描目录，返回文件树结构。忽略以 '.' 开头的隐藏文件和目录。"""
        file_count = 0
        dir_count = 0

        def _build_node(abs_path: str, rel_path: str) -> FileNode:
            nonlocal file_count, dir_count
            name = os.path.basename(abs_path)
            is_dir = os.path.isdir(abs_path)

            if is_dir:
                dir_count += 1
                children_nodes: List[FileNode] = []
                try:
                    entries = os.listdir(abs_path)
                except PermissionError:
                    entries = []
                for entry_name in entries:
                    if self._should_ignore(entry_name):
                        continue
                    entry_abs = os.path.join(abs_path, entry_name)
                    entry_rel = os.path.join(rel_path, entry_name) if rel_path else entry_name
                    children_nodes.append(_build_node(entry_abs, entry_rel))
                # Sort: directories first, then files, both alphabetically
                children_nodes.sort(key=lambda n: (not n.is_dir, n.name.lower()))
                return FileNode(
                    name=name,
                    path=rel_path,
                    is_dir=True,
                    extension='',
                    children=children_nodes,
                    display_name=name,
                    page_path='',
                )
            else:
                file_count += 1
                _, ext = os.path.splitext(name)
                extension = ext.lower()
                display_name = os.path.splitext(name)[0]
                # Build page_path: pages/{rel_dir}/{filename_no_ext}.html
                # Spaces in filename replaced with underscores
                safe_name = display_name.replace(' ', '_') + '.html'
                rel_dir = os.path.dirname(rel_path)
                if rel_dir:
                    page_path = os.path.join('pages', rel_dir, safe_name)
                else:
                    page_path = os.path.join('pages', safe_name)
                # Normalize path separators to forward slashes
                page_path = page_path.replace('\\', '/')
                return FileNode(
                    name=name,
                    path=rel_path,
                    is_dir=False,
                    extension=extension,
                    children=[],
                    display_name=display_name,
                    page_path=page_path,
                )

        root_node = _build_node(root_path, '')
        return FileTree(root=root_node, file_count=file_count, dir_count=dir_count)

    def _should_ignore(self, name: str) -> bool:
        """判断文件或目录是否应被忽略（以 '.' 开头的隐藏项）。"""
        return name.startswith('.')


class ContentRenderer:
    """根据文件类型渲染文件内容，返回 HTML 片段。"""

    def render(self, file_path: str, file_type: str) -> str:
        """根据文件类型渲染文件内容，返回 HTML 片段（body 部分）。"""
        renderer_name = RENDERER_MAP.get(file_type)
        if renderer_name:
            method = getattr(self, renderer_name)
            return method(file_path)
        # 不在映射中：尝试纯文本渲染
        try:
            return self._render_plaintext(file_path)
        except Exception:
            return self._render_unsupported()

    def _render_markdown(self, path: str) -> str:
        """将 Markdown 转换为 HTML。"""
        import markdown
        with open(path, 'r', encoding='utf-8') as f:
            text = f.read()

        # 将本地图片引用转换为 base64 data URI
        md_dir = os.path.dirname(path)
        def _replace_img(match):
            alt = match.group(1)
            img_path = match.group(2)
            if img_path.startswith(('http://', 'https://', 'data:')):
                return match.group(0)
            abs_img = os.path.join(md_dir, img_path)
            if os.path.isfile(abs_img):
                mime, _ = mimetypes.guess_type(abs_img)
                if not mime:
                    mime = 'image/png'
                with open(abs_img, 'rb') as img_f:
                    b64 = base64.b64encode(img_f.read()).decode('ascii')
                return f'![{alt}](data:{mime};base64,{b64})'
            return match.group(0)
        text = re.sub(r'!\[([^\]]*)\]\(([^)]+)\)', _replace_img, text)

        # 处理 Mermaid 代码块：替换为 div 以便前端渲染
        def _replace_mermaid(match):
            code = html_module.escape(match.group(1))
            return f'<div class="mermaid">{code}</div>'
        text = re.sub(r'```mermaid\s*\n(.*?)```', _replace_mermaid, text, flags=re.DOTALL)

        # 使用 markdown + pymdown-extensions 渲染
        extensions = [
            'tables', 'fenced_code', 'codehilite', 'toc',
            'pymdownx.superfences', 'pymdownx.arithmatex',
        ]
        extension_configs = {
            'codehilite': {'css_class': 'code-highlight', 'guess_lang': True},
            'pymdownx.arithmatex': {'generic': True},
            'toc': {'permalink': False, 'slugify': lambda value, separator: re.sub(r'-+', '-', re.sub(r'[^\w\u4e00-\u9fff-]', '', value.lower().replace(' ', separator))).strip('-')},
        }
        html_content = markdown.markdown(
            text, extensions=extensions, extension_configs=extension_configs,
        )

        # 将 arithmatex 输出的 LaTeX 公式转换为 MathML（纯静态，无需 JS）
        html_content = self._convert_arithmatex_to_mathml(html_content)

        # 确保外部链接在新标签页打开
        html_content = re.sub(
            r'<a\s+href="(https?://[^"]+)"',
            r'<a href="\1" target="_blank" rel="noopener noreferrer"',
            html_content,
        )

        return f'<div class="markdown-body">{html_content}</div>'

    @staticmethod
    def _convert_arithmatex_to_mathml(html_content: str) -> str:
        """将 arithmatex 生成的 LaTeX 分隔符转换为 MathML，实现纯静态渲染。"""
        import latex2mathml.converter

        def _latex_to_mathml(latex_str: str, display: bool = False) -> str:
            try:
                mathml = latex2mathml.converter.convert(latex_str)
                if display:
                    mathml = mathml.replace('<math', '<math display="block"', 1)
                return mathml
            except Exception:
                # 转换失败时保留原始 LaTeX 文本
                escaped = html_module.escape(latex_str)
                if display:
                    return f'<div class="math-display"><code>{escaped}</code></div>'
                return f'<code>{escaped}</code>'

        # 块级公式: <div class="arithmatex">\[...\]</div>
        def _replace_block(match):
            latex = match.group(1).strip()
            return _latex_to_mathml(latex, display=True)
        html_content = re.sub(
            r'<div class="arithmatex">\s*\\\[(.*?)\\\]\s*</div>',
            _replace_block, html_content, flags=re.DOTALL,
        )

        # 行内公式: <span class="arithmatex">\(...\)</span>
        def _replace_inline(match):
            latex = match.group(1).strip()
            return _latex_to_mathml(latex, display=False)
        html_content = re.sub(
            r'<span class="arithmatex">\s*\\\((.*?)\\\)\s*</span>',
            _replace_inline, html_content, flags=re.DOTALL,
        )

        return html_content

    def _render_pdf(self, path: str) -> str:
        """将 PDF 渲染为 canvas，base64 存在隐藏 textarea 中避免 JS 阻塞。"""
        with open(path, 'rb') as f:
            b64 = base64.b64encode(f.read()).decode('ascii')
        # 将 base64 分块存入多个隐藏 textarea，每块 512KB，避免单个 DOM 节点过大
        chunk_size = 512 * 1024
        chunks = [b64[i:i+chunk_size] for i in range(0, len(b64), chunk_size)]
        textarea_html = ''
        for i, chunk in enumerate(chunks):
            textarea_html += f'<textarea id="pdf-chunk-{i}" style="display:none;">{chunk}</textarea>\n'
        return (
            f'{textarea_html}'
            f'<input type="hidden" id="pdf-chunk-count" value="{len(chunks)}">\n'
            '<style>\n'
            '.pdf-page-wrap{position:relative;margin-bottom:2px;}\n'
            '.pdf-page-wrap canvas{width:100%;height:auto;display:block;}\n'
            '.pdf-page-loading{text-align:center;padding:30px 0;color:#888;font-size:0.9em;background:#f9f9f9;border:1px solid #eee;border-radius:4px;}\n'
            '</style>\n'
            '<div id="pdf-progress" style="text-align:center;padding:40px;color:#555;font-size:0.95em;">\n'
            '  <div id="pdf-status">准备加载 PDF...</div>\n'
            '  <div style="width:60%;max-width:400px;margin:12px auto;background:#eee;border-radius:4px;height:6px;">\n'
            '    <div id="pdf-bar" style="width:0%;height:100%;background:#4a90d9;border-radius:4px;transition:width 0.2s;"></div>\n'
            '  </div>\n'
            '</div>\n'
            '<div id="pdf-container" style="width:100%;max-width:900px;margin:0 auto;display:none;"></div>\n'
            '<script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>\n'
            '<script>\n'
            'pdfjsLib.GlobalWorkerOptions.workerSrc="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.worker.min.js";\n'
            'var progress=document.getElementById("pdf-progress");\n'
            'var statusEl=document.getElementById("pdf-status");\n'
            'var bar=document.getElementById("pdf-bar");\n'
            'var container=document.getElementById("pdf-container");\n'
            'function setProgress(pct,msg){bar.style.width=pct+"%";if(msg)statusEl.textContent=msg;}\n'
            '\n'
            'requestAnimationFrame(function(){setProgress(5,"正在解码 PDF 数据...");requestAnimationFrame(function(){\n'
            '  var count=parseInt(document.getElementById("pdf-chunk-count").value);\n'
            '  var b64="";\n'
            '  for(var i=0;i<count;i++){\n'
            '    var el=document.getElementById("pdf-chunk-"+i);\n'
            '    b64+=el.value;\n'
            '    el.remove();\n'
            '  }\n'
            '  setProgress(10,"正在转换数据...");\n'
            '  var bin=atob(b64);b64=null;\n'
            '  var len=bin.length,bytes=new Uint8Array(len);\n'
            '  for(var j=0;j<len;j++)bytes[j]=bin.charCodeAt(j);\n'
            '  bin=null;\n'
            '  setProgress(20,"正在解析 PDF 文档...");\n'
            '  var pdfDoc=null,rendered={},totalPages=0;\n'
            '  var dpr=window.devicePixelRatio||1;\n'
            '  var RENDER_SCALE=dpr<2?2:Math.min(dpr,3);\n'
            '  function renderPage(num){\n'
            '    if(rendered[num])return Promise.resolve();\n'
            '    rendered[num]=true;\n'
            '    var hint=document.getElementById("pdf-hint-"+num);\n'
            '    if(hint)hint.textContent="正在渲染第 "+num+"/"+totalPages+" 页...";\n'
            '    return pdfDoc.getPage(num).then(function(page){\n'
            '      var vw=container.clientWidth||document.documentElement.clientWidth-40;\n'
            '      var baseScale=vw/page.getViewport({scale:1}).width;\n'
            '      var viewport=page.getViewport({scale:baseScale*RENDER_SCALE});\n'
            '      var canvas=document.getElementById("pdf-page-"+num);\n'
            '      if(!canvas)return;\n'
            '      canvas.width=viewport.width;\n'
            '      canvas.height=viewport.height;\n'
            '      return page.render({canvasContext:canvas.getContext("2d"),viewport:viewport}).promise.then(function(){if(hint)hint.remove();});\n'
            '    });\n'
            '  }\n'
            '  pdfjsLib.getDocument({data:bytes}).promise.then(function(pdf){\n'
            '    pdfDoc=pdf;totalPages=pdf.numPages;\n'
            '    setProgress(25,"共 "+totalPages+" 页，正在渲染...");\n'
            '    container.style.display="";\n'
            '    for(var p=1;p<=pdf.numPages;p++){\n'
            '      var wrap=document.createElement("div");wrap.className="pdf-page-wrap";\n'
            '      var canvas=document.createElement("canvas");canvas.id="pdf-page-"+p;\n'
            '      var hint=document.createElement("div");hint.id="pdf-hint-"+p;\n'
            '      hint.className="pdf-page-loading";hint.textContent="第 "+p+" 页 等待加载...";\n'
            '      wrap.appendChild(hint);wrap.appendChild(canvas);container.appendChild(wrap);\n'
            '    }\n'
            '    renderPage(1).then(function(){\n'
            '      progress.remove();\n'
            '      if(pdf.numPages>1)renderPage(2);\n'
            '      var observer=new IntersectionObserver(function(entries){\n'
            '        entries.forEach(function(e){\n'
            '          if(e.isIntersecting){\n'
            '            var num=parseInt(e.target.id.replace("pdf-page-",""));\n'
            '            renderPage(num);if(num+1<=pdf.numPages)renderPage(num+1);\n'
            '            observer.unobserve(e.target);\n'
            '          }\n'
            '        });\n'
            '      },{rootMargin:"200px"});\n'
            '      for(var p=3;p<=pdf.numPages;p++)observer.observe(document.getElementById("pdf-page-"+p));\n'
            '    });\n'
            '  });\n'
            '});});\n'
            '</script>'
        )

    def _render_html(self, path: str) -> str:
        """将 HTML 文件内容内嵌渲染。"""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        # 内嵌图片和样式资源为 base64
        html_dir = os.path.dirname(path)
        def _inline_resource(match):
            attr = match.group(1)  # src= or href=
            url = match.group(2)
            if url.startswith(('http://', 'https://', 'data:')):
                return match.group(0)
            res_path = os.path.join(html_dir, url)
            if os.path.isfile(res_path):
                mime, _ = mimetypes.guess_type(res_path)
                if not mime:
                    mime = 'application/octet-stream'
                with open(res_path, 'rb') as rf:
                    b64 = base64.b64encode(rf.read()).decode('ascii')
                return f'{attr}"data:{mime};base64,{b64}"'
            return match.group(0)
        content = re.sub(r'(src=|href=)"([^"]+)"', _inline_resource, content)
        escaped = content.replace('&', '&amp;').replace('"', '&quot;')
        return (f'<div class="html-viewer" style="position:relative;">'
                f'<div id="html-loading" style="text-align:center;color:#888;padding:20px;">加载 HTML 内容中...</div>'
                f'<iframe srcdoc="{escaped}" '
                f'width="100%" height="800px" style="border:none;" '
                f'sandbox="allow-scripts allow-same-origin allow-popups allow-top-navigation-by-user-activation" '
                f'onload="document.getElementById(\'html-loading\').remove();"></iframe>'
                f'</div>')

    def _render_image(self, path: str) -> str:
        """将图片转换为 base64 data URI，带加载占位。"""
        mime, _ = mimetypes.guess_type(path)
        if not mime:
            ext = os.path.splitext(path)[1].lower()
            mime_map = {'.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png',
                        '.gif': 'image/gif', '.svg': 'image/svg+xml', '.bmp': 'image/bmp',
                        '.webp': 'image/webp'}
            mime = mime_map.get(ext, 'image/png')
        fsize = os.path.getsize(path)
        size_str = f'{fsize/1024:.0f} KB' if fsize < 1024*1024 else f'{fsize/1024/1024:.1f} MB'
        with open(path, 'rb') as f:
            b64 = base64.b64encode(f.read()).decode('ascii')
        return (f'<div class="image-viewer" style="text-align:center;">'
                f'<div id="img-loading" style="color:#888;padding:20px;">加载图片中 ({size_str})...</div>'
                f'<img src="data:{mime};base64,{b64}" '
                f'style="max-width:100%;height:auto;display:none;" alt="image" '
                f'onload="this.style.display=\'block\';document.getElementById(\'img-loading\').remove();" />'
                f'</div>')

    def _render_docx(self, path: str) -> str:
        """提取 DOCX 文档内容并转换为 HTML。"""
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        try:
            doc = Document(path)
        except Exception as e:
            return f'<div class="error"><p>无法解析 DOCX 文件: {html_module.escape(str(e))}</p></div>'

        # 收集超链接关系映射 rId -> URL
        hyperlink_rels = {}
        for rel_id, rel in doc.part.rels.items():
            if "hyperlink" in rel.reltype:
                hyperlink_rels[rel_id] = rel._target

        parts = []
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else ''
            text = ''
            # 遍历段落的 XML 子元素以正确处理超链接
            from lxml import etree
            nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                     'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'}
            for child in para._element:
                tag = etree.QName(child.tag).localname if isinstance(child.tag, str) else ''
                if tag == 'hyperlink':
                    r_id = child.get(f'{{{nsmap["r"]}}}id', '')
                    url = hyperlink_rels.get(r_id, '')
                    link_text = ''
                    for r in child.findall('.//w:t', nsmap):
                        link_text += r.text or ''
                    if url:
                        escaped_url = html_module.escape(url)
                        text += f'<a href="{escaped_url}" target="_blank" rel="noopener noreferrer">{html_module.escape(link_text)}</a>'
                    else:
                        text += html_module.escape(link_text)
                elif tag == 'r':
                    run_text = ''
                    for t_elem in child.findall('.//w:t', nsmap):
                        run_text += t_elem.text or ''
                    t = html_module.escape(run_text)
                    # 检查粗体/斜体
                    rpr = child.find('w:rPr', nsmap)
                    if rpr is not None:
                        if rpr.find('w:b', nsmap) is not None:
                            t = f'<strong>{t}</strong>'
                        if rpr.find('w:i', nsmap) is not None:
                            t = f'<em>{t}</em>'
                    text += t

            if not text.strip():
                continue
            if 'Heading 1' in style_name:
                parts.append(f'<h1>{text}</h1>')
            elif 'Heading 2' in style_name:
                parts.append(f'<h2>{text}</h2>')
            elif 'Heading 3' in style_name:
                parts.append(f'<h3>{text}</h3>')
            elif 'List' in style_name:
                parts.append(f'<li>{text}</li>')
            else:
                parts.append(f'<p>{text}</p>')

        # 表格
        for table in doc.tables:
            parts.append('<table class="docx-table">')
            for i, row in enumerate(table.rows):
                parts.append('<tr>')
                tag = 'th' if i == 0 else 'td'
                for cell in row.cells:
                    parts.append(f'<{tag}>{html_module.escape(cell.text)}</{tag}>')
                parts.append('</tr>')
            parts.append('</table>')

        # 图片（内嵌 base64）
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                try:
                    img_data = rel.target_part.blob
                    content_type = rel.target_part.content_type
                    b64 = base64.b64encode(img_data).decode('ascii')
                    parts.append(f'<img src="data:{content_type};base64,{b64}" '
                                 f'style="max-width:100%;height:auto;" />')
                except Exception:
                    pass

        return (f'<div class="docx-content">'
                f'<div style="color:#888;font-size:0.85em;margin-bottom:12px;">'
                f'共 {len(doc.paragraphs)} 个段落，{len(doc.tables)} 个表格</div>'
                f'{"".join(parts)}</div>')

    def _render_xlsx(self, path: str) -> str:
        """将 XLSX 电子表格转换为 HTML 表格。"""
        from openpyxl import load_workbook
        try:
            wb = load_workbook(path, read_only=True, data_only=True)
        except Exception as e:
            return f'<div class="error"><p>无法解析 XLSX 文件: {html_module.escape(str(e))}</p></div>'

        sheets = wb.sheetnames
        parts = []

        # 标签页切换按钮
        if len(sheets) > 1:
            parts.append('<div class="xlsx-tabs">')
            for i, name in enumerate(sheets):
                active = ' active' if i == 0 else ''
                parts.append(f'<button class="xlsx-tab{active}" '
                             f'onclick="document.querySelectorAll(\'.xlsx-sheet\').forEach(s=>s.style.display=\'none\');'
                             f'document.getElementById(\'sheet-{i}\').style.display=\'block\';'
                             f'document.querySelectorAll(\'.xlsx-tab\').forEach(b=>b.classList.remove(\'active\'));'
                             f'this.classList.add(\'active\');">'
                             f'{html_module.escape(name)}</button>')
            parts.append('</div>')

        for i, name in enumerate(sheets):
            ws = wb[name]
            display = 'block' if i == 0 else 'none'
            row_count = 0
            parts.append(f'<div id="sheet-{i}" class="xlsx-sheet" style="display:{display};">')
            parts.append('<div style="overflow-x:auto;">')
            parts.append('<table class="xlsx-table">')
            for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                row_count += 1
                parts.append('<tr>')
                for cell in row:
                    tag = 'th' if row_idx == 0 else 'td'
                    val = html_module.escape(str(cell)) if cell is not None else ''
                    parts.append(f'<{tag}>{val}</{tag}>')
                parts.append('</tr>')
            parts.append('</table></div>')
            parts.append(f'<div style="color:#888;font-size:0.8em;margin-top:4px;">{row_count} 行</div>')
            parts.append('</div>')

        wb.close()
        return '\n'.join(parts)

    def _render_csv(self, path: str) -> str:
        """将 CSV 数据转换为 HTML 表格。"""
        content = None
        for enc in ['utf-8', 'gbk', 'latin-1']:
            try:
                with open(path, 'r', encoding=enc, newline='') as f:
                    content = f.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        if content is None:
            return self._render_unsupported()

        reader = csv.reader(content.splitlines())
        rows = list(reader)
        if not rows:
            return '<p>空的 CSV 文件</p>'

        html_parts = ['<div style="overflow-x:auto;">', '<table class="csv-table">']
        for i, row in enumerate(rows):
            if i == 0:
                html_parts.append('<thead><tr>')
                for cell in row:
                    html_parts.append(f'<th>{html_module.escape(cell)}</th>')
                html_parts.append('</tr></thead><tbody>')
            else:
                cls = 'even' if i % 2 == 0 else 'odd'
                html_parts.append(f'<tr class="{cls}">')
                for cell in row:
                    html_parts.append(f'<td>{html_module.escape(cell)}</td>')
                html_parts.append('</tr>')
        html_parts.append('</tbody></table></div>')
        html_parts.append(f'<div style="color:#888;font-size:0.8em;margin-top:4px;">{len(rows)} 行 × {len(rows[0]) if rows else 0} 列</div>')
        return '\n'.join(html_parts)

    def _render_json(self, path: str) -> str:
        """将 JSON 数据转换为语法高亮的树形结构。"""
        with open(path, 'r', encoding='utf-8') as f:
            raw = f.read()
        try:
            data = json.loads(raw)
            formatted = json.dumps(data, indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            # 无效 JSON，以纯文本代码块展示
            escaped = html_module.escape(raw)
            return f'<pre><code>{escaped}</code></pre>'

        def _json_to_html(obj, depth=0):
            """递归生成可折叠的 JSON HTML 树。"""
            indent = '  ' * depth
            if isinstance(obj, dict):
                if not obj:
                    return '<span class="json-brace">{}</span>'
                items = []
                for k, v in obj.items():
                    key_html = f'<span class="json-key">"{html_module.escape(str(k))}"</span>'
                    val_html = _json_to_html(v, depth + 1)
                    items.append(f'{key_html}: {val_html}')
                inner = ',\n'.join(items)
                return (f'<details open><summary class="json-brace">{{</summary>'
                        f'<div class="json-indent">{inner}</div>'
                        f'<span class="json-brace">}}</span></details>')
            elif isinstance(obj, list):
                if not obj:
                    return '<span class="json-bracket">[]</span>'
                items = [_json_to_html(item, depth + 1) for item in obj]
                inner = ',\n'.join(items)
                return (f'<details open><summary class="json-bracket">[</summary>'
                        f'<div class="json-indent">{inner}</div>'
                        f'<span class="json-bracket">]</span></details>')
            elif isinstance(obj, str):
                escaped_str = html_module.escape(obj)
                if re.match(r'https?://', obj):
                    return f'<span class="json-string">"<a href="{escaped_str}" target="_blank" rel="noopener noreferrer">{escaped_str}</a>"</span>'
                return f'<span class="json-string">"{escaped_str}"</span>'
            elif isinstance(obj, bool):
                return f'<span class="json-bool">{str(obj).lower()}</span>'
            elif obj is None:
                return '<span class="json-null">null</span>'
            else:
                return f'<span class="json-number">{html_module.escape(str(obj))}</span>'

        tree_html = _json_to_html(data)
        return f'<div class="json-viewer">{tree_html}</div>'

    def _render_mermaid(self, path: str) -> str:
        """将 Mermaid 语法内嵌到页面，带渲染状态提示。"""
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        escaped = html_module.escape(content)
        return (f'<div id="mermaid-loading" style="text-align:center;color:#888;padding:20px;">正在渲染 Mermaid 图表...</div>'
                f'<div class="mermaid" style="display:none;" '
                f'onrender="this.style.display=\'block\';document.getElementById(\'mermaid-loading\').remove();">'
                f'{escaped}</div>'
                f'<script>document.addEventListener("DOMContentLoaded",function(){{'
                f'setTimeout(function(){{document.querySelector(".mermaid").style.display="block";'
                f'var l=document.getElementById("mermaid-loading");if(l)l.remove();}},2000);'
                f'}});</script>')

    def _render_latex(self, path: str) -> str:
        """将 LaTeX 源码渲染为 MathML。"""
        import latex2mathml.converter
        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()
        escaped_code = html_module.escape(content)
        parts = []
        parts.append('<div class="latex-content">')
        parts.append('<h3>源码</h3>')
        parts.append(f'<pre><code class="language-latex">{escaped_code}</code></pre>')
        parts.append('<h3>渲染结果</h3>')
        try:
            mathml = latex2mathml.converter.convert(content)
            parts.append(f'<div class="math-display">{mathml}</div>')
        except Exception:
            parts.append(f'<div class="katex-render">{escaped_code}</div>')
        parts.append('</div>')
        return '\n'.join(parts)

    def _render_code(self, path: str) -> str:
        """使用 Pygments 对代码文件进行语法高亮渲染。"""
        from pygments import highlight
        from pygments.lexers import get_lexer_for_filename, TextLexer
        from pygments.formatters import HtmlFormatter

        with open(path, 'r', encoding='utf-8') as f:
            content = f.read()

        line_count = content.count('\n') + (1 if content and not content.endswith('\n') else 0)
        fsize = os.path.getsize(path)
        size_str = f'{fsize/1024:.1f} KB' if fsize >= 1024 else f'{fsize} B'

        try:
            lexer = get_lexer_for_filename(path)
            lang = lexer.name
        except Exception:
            lexer = TextLexer()
            lang = 'Text'

        formatter = HtmlFormatter(linenos=True, cssclass='code-highlight')
        code_html = highlight(content, lexer, formatter)
        info = f'<div style="color:#888;font-size:0.8em;margin-bottom:8px;">{lang} · {line_count} 行 · {size_str}</div>'
        return info + code_html

    def _render_plaintext(self, path: str) -> str:
        """以纯文本形式渲染文件内容。"""
        encodings = ['utf-8', 'gbk', 'latin-1']
        for enc in encodings:
            try:
                with open(path, 'r', encoding=enc) as f:
                    content = f.read()
                escaped = html_module.escape(content)
                return f'<pre><code>{escaped}</code></pre>'
            except (UnicodeDecodeError, UnicodeError):
                continue
        raise ValueError(f"Cannot read file as text: {path}")

    def _render_unsupported(self) -> str:
        """显示不支持预览的提示信息。"""
        return '<div class="unsupported"><p>该文件格式暂不支持预览</p></div>'


class PageGenerator:
    """将渲染后的 HTML 片段包装为完整的 Document_Page HTML 文件。"""

    def generate_page(self, title: str, body_html: str, page_type: str) -> str:
        """将 HTML 片段包装为完整的 Document_Page。"""
        extra_head = ''
        extra_body = ''
        if page_type in ('.md', '.tex', '.latex'):
            extra_head += self._get_katex_head()
        if page_type in ('.mermaid', '.md'):
            extra_body += f'<script>{self._get_mermaid_js()}</script>\n'
            extra_body += '<script>document.addEventListener("DOMContentLoaded",function(){mermaid.initialize({startOnLoad:true});});</script>\n'
        if page_type in RENDERER_MAP and RENDERER_MAP.get(page_type) == '_render_code':
            extra_head += f'<style>{self._get_pygments_css()}</style>\n'
        if page_type == '.md':
            extra_head += f'<style>{self._get_pygments_css()}</style>\n'
        return self._wrap_with_template(title, body_html + extra_body, extra_head)

    def _wrap_with_template(self, title: str, body: str, extra_head: str) -> str:
        """使用 HTML 模板包装内容。"""
        return f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{html_module.escape(title)}</title>
<style>
html {{ scroll-behavior: smooth; }}
body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
       margin: 0; padding: 20px 40px; line-height: 1.6; color: #333; background: #fff; }}
a {{ color: #0366d6; text-decoration: none; }}
a:hover {{ text-decoration: underline; }}
img {{ max-width: 100%; height: auto; }}
table {{ border-collapse: collapse; width: 100%; margin: 1em 0; overflow-x: auto; display: block; }}
th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; white-space: nowrap; }}
th {{ background: #f5f5f5; font-weight: 600; }}
tr:nth-child(even) {{ background: #fafafa; }}
pre {{ background: #f6f8fa; padding: 16px; border-radius: 6px; overflow-x: auto; font-size: 0.85em; }}
code {{ font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace; font-size: 0.9em; }}
.unsupported {{ text-align: center; padding: 60px 20px; color: #999; }}
.markdown-body {{ max-width: 900px; margin: 0 auto; }}
@media (max-width: 600px) {{
  body {{ padding: 10px 12px; font-size: 0.95em; }}
  table {{ font-size: 0.85em; }}
  th, td {{ padding: 4px 6px; }}
  pre {{ padding: 10px; font-size: 0.8em; }}
  .markdown-body {{ padding: 0; }}
}}
.json-viewer {{ font-family: monospace; font-size: 0.9em; }}
.json-indent {{ margin-left: 20px; }}
.json-key {{ color: #881391; }}
.json-string {{ color: #1a1aa6; }}
.json-number {{ color: #1c6b48; }}
.json-bool {{ color: #d63200; }}
.json-null {{ color: #808080; }}
.csv-table th {{ background: #4a90d9; color: white; }}
.csv-table tr.odd {{ background: #f9f9f9; }}
.xlsx-tabs {{ margin-bottom: 10px; }}
.xlsx-tab {{ padding: 8px 16px; border: 1px solid #ddd; background: #f5f5f5; cursor: pointer; border-bottom: none; }}
.xlsx-tab.active {{ background: white; font-weight: bold; }}
.docx-table th {{ background: #f0f0f0; }}
.katex-render {{ font-size: 1.2em; line-height: 2; }}
.mermaid {{ text-align: center; margin: 20px 0; }}
#page-loading {{ position: fixed; top: 0; left: 0; width: 100%; height: 100%; background: #fff;
  display: flex; align-items: center; justify-content: center; z-index: 9999;
  flex-direction: column; color: #888; font-size: 0.95em; }}
#page-loading .bar-wrap {{ width: 60%; max-width: 300px; height: 4px; background: #eee;
  border-radius: 2px; margin-top: 12px; overflow: hidden; }}
#page-loading .bar {{ width: 30%; height: 100%; background: #4a90d9; border-radius: 2px;
  animation: loading-slide 1.2s ease-in-out infinite; }}
@keyframes loading-slide {{ 0% {{ transform: translateX(-100%); }} 100% {{ transform: translateX(400%); }} }}
</style>
{extra_head}
</head>
<body oncontextmenu="return false;">
<div id="page-loading"><span>加载中...</span><div class="bar-wrap"><div class="bar"></div></div></div>
{body}
<script>document.getElementById("page-loading").remove();</script>
</body>
</html>'''

    def _get_katex_head(self) -> str:
        """获取数学公式相关的 CSS 样式。"""
        return '<style>.arithmatex { font-size: 1.1em; } .math-display { overflow-x: auto; margin: 1em 0; text-align: center; }</style>'

    def _get_mermaid_js(self) -> str:
        """获取 Mermaid.js 脚本。"""
        cache_dir = os.path.join('.cache')
        cache_file = os.path.join(cache_dir, 'mermaid.min.js')
        if os.path.isfile(cache_file):
            with open(cache_file, 'r', encoding='utf-8') as f:
                return f.read()
        # 尝试从 CDN 下载并缓存
        try:
            import urllib.request
            url = 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.min.js'
            os.makedirs(cache_dir, exist_ok=True)
            urllib.request.urlretrieve(url, cache_file)
            with open(cache_file, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception:
            return '/* Mermaid.js not available - render as text */'

    def _get_pygments_css(self) -> str:
        """获取 Pygments CSS 样式。"""
        from pygments.formatters import HtmlFormatter
        return HtmlFormatter(cssclass='code-highlight').get_style_defs()


class IndexGenerator:
    """根据文件树生成轻量的 index.html 内容。"""

    def generate_index(self, file_tree: FileTree) -> str:
        """根据文件树生成 index.html 内容。"""
        tree_html = ''
        for child in file_tree.root.children:
            tree_html += self._build_tree_html(child)
        return f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>文档索引</title>
<style>
{self._build_css()}
</style>
</head>
<body oncontextmenu="return false;">
<div class="container">
<h1 class="title">📚 文档索引</h1>
<p class="stats">{file_tree.file_count} 个文件，{file_tree.dir_count} 个目录</p>
<div class="tree-container">
<ul class="tree-root">
{tree_html}
</ul>
</div>
</div>
<script>
{self._build_js()}
</script>
</body>
</html>'''

    def _build_tree_html(self, node: FileNode) -> str:
        """递归生成目录树 HTML。"""
        if node.is_dir:
            children_html = ''
            for child in node.children:
                children_html += self._build_tree_html(child)
            return (f'<li class="tree-dir">'
                    f'<span class="dir-toggle collapsed" onclick="toggleDir(this)">▶</span>'
                    f'<span class="dir-name" onclick="toggleDir(this.previousElementSibling)">📁 {html_module.escape(node.display_name)}</span>'
                    f'<ul class="tree-children" style="display:none;">{children_html}</ul>'
                    f'</li>\n')
        else:
            icon = self._get_file_icon(node.extension)
            return (f'<li class="tree-file">'
                    f'<a href="{html_module.escape(node.page_path)}" target="_blank">'
                    f'{icon} {html_module.escape(node.display_name)}</a>'
                    f'</li>\n')

    def _get_file_icon(self, file_type: str) -> str:
        """为不同文件类型返回对应的图标标识。"""
        icon_map = {
            '.md': '📝', '.pdf': '📕', '.html': '🌐', '.htm': '🌐',
            '.jpg': '🖼️', '.jpeg': '🖼️', '.png': '🖼️', '.gif': '🖼️',
            '.svg': '🖼️', '.bmp': '🖼️', '.webp': '🖼️',
            '.docx': '📄', '.xlsx': '📊',
            '.csv': '📊', '.json': '📋',
            '.mermaid': '📐', '.tex': '📐', '.latex': '📐',
            '.py': '🐍', '.java': '☕', '.js': '📜', '.ts': '📜',
            '.cpp': '⚙️', '.hpp': '⚙️', '.c': '⚙️', '.h': '⚙️',
            '.go': '🔵', '.rs': '🦀', '.rb': '💎', '.php': '🐘',
            '.swift': '🍎', '.kt': '🟣', '.scala': '🔴',
            '.sh': '🖥️', '.bash': '🖥️', '.css': '🎨', '.sql': '🗃️',
        }
        return icon_map.get(file_type, '📄')

    def _build_css(self) -> str:
        """生成索引页的内嵌 CSS。"""
        return '''
* { margin: 0; padding: 0; box-sizing: border-box; }
body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", sans-serif;
       background: #f8f9fa; color: #333; min-height: 100vh; }
.container { max-width: 800px; margin: 0 auto; padding: 40px 20px; }
.title { font-size: 1.8em; margin-bottom: 4px; color: #2c3e50; }
.stats { color: #888; font-size: 0.9em; margin-bottom: 24px; }
.tree-container { background: white; border-radius: 8px; padding: 20px 24px;
                  box-shadow: 0 1px 3px rgba(0,0,0,0.1); }
ul { list-style: none; }
.tree-root { padding: 0; }
.tree-children { padding-left: 20px; }
.tree-dir { margin: 4px 0; }
.tree-file { margin: 2px 0; padding: 2px 0; }
.dir-toggle { cursor: pointer; user-select: none; display: inline-block; width: 16px;
              font-size: 0.7em; color: #888; }
.dir-name { font-weight: 500; cursor: pointer; }
.tree-file a { text-decoration: none; color: #0366d6; transition: color 0.15s; }
.tree-file a:hover { color: #0250a3; text-decoration: underline; }
@media (max-width: 600px) {
  .container { padding: 20px 12px; }
  .title { font-size: 1.4em; }
  .tree-container { padding: 12px 14px; }
  .tree-children { padding-left: 14px; }
  .tree-file a { font-size: 0.95em; }
}
'''

    def _build_js(self) -> str:
        """生成索引页的内嵌 JavaScript。"""
        return '''
function toggleDir(el) {
    var children = el.parentElement.querySelector('.tree-children');
    if (children.style.display === 'none') {
        children.style.display = '';
        el.classList.remove('collapsed');
        el.textContent = '▼';
    } else {
        children.style.display = 'none';
        el.classList.add('collapsed');
        el.textContent = '▶';
    }
}
'''


class IndexBuilder:
    """主入口：串联完整的生成流程。"""

    MAX_FILE_SIZE = 0  # 不限制文件大小

    def build(self, docs_dir: str = 'docs', output_dir: str = '.') -> None:
        """主入口：扫描 → 渲染 → 生成 → 写入。"""
        import sys
        if not os.path.isdir(docs_dir):
            print(f"错误: 目录 '{docs_dir}' 不存在", file=sys.stderr)
            sys.exit(1)

        scanner = DirectoryScanner()
        renderer = ContentRenderer()
        page_gen = PageGenerator()
        index_gen = IndexGenerator()

        # 1. 扫描
        print("📂 扫描文档目录...")
        file_tree = scanner.scan(docs_dir)
        print(f"   找到 {file_tree.file_count} 个文件，{file_tree.dir_count} 个目录")

        # 2. 清空 pages/
        pages_dir = os.path.join(output_dir, 'pages')
        self._clean_pages_dir(pages_dir)

        # 3. 渲染并生成 Document_Page
        print("📝 渲染页面...")
        stats = BuildStats(
            total_files=file_tree.file_count,
            total_dirs=file_tree.dir_count,
        )
        pages = {}  # page_path -> html_content

        def _process_node(node: FileNode):
            if node.is_dir:
                for child in node.children:
                    _process_node(child)
                return

            abs_path = os.path.join(docs_dir, node.path)
            try:
                os.path.getsize(abs_path)
            except OSError:
                stats.errors.append(f"无法访问: {node.path}")
                print(f"  ✗ {node.path} (无法访问)")
                return

            try:
                processed = stats.total_pages + len(stats.errors) + 1
                print(f"  [{processed}/{file_tree.file_count}] 渲染 {node.path} ...", end='', flush=True)
                body_html = renderer.render(abs_path, node.extension)
                full_html = page_gen.generate_page(node.display_name, body_html, node.extension)
                # page_path 去掉 'pages/' 前缀得到相对于 pages/ 的路径
                rel_page = node.page_path
                if rel_page.startswith('pages/'):
                    rel_page = rel_page[6:]
                pages[rel_page] = full_html
                stats.total_pages += 1
                stats.file_type_counts[node.extension] = stats.file_type_counts.get(node.extension, 0) + 1
                print(f" ✓")
            except Exception as e:
                stats.errors.append(f"渲染失败 {node.path}: {str(e)}")
                print(f" ✗ ({e})")

        _process_node(file_tree.root)

        # 4. 生成 index.html
        print("📋 生成 index.html...")
        index_html = index_gen.generate_index(file_tree)

        # 5. 写入文件
        print("💾 写入文件...")
        index_path = os.path.join(output_dir, 'index.html')
        self._write_index(index_html, index_path)
        stats.index_size_bytes = len(index_html.encode('utf-8'))

        self._write_pages(pages, pages_dir)
        stats.pages_size_bytes = sum(len(v.encode('utf-8')) for v in pages.values())

        # 6. 输出摘要
        self._print_summary(stats)

    def _write_index(self, content: str, output_path: str) -> None:
        """将 index.html 写入指定路径。"""
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)

    def _write_pages(self, pages: dict, pages_dir: str) -> None:
        """将各 Document_Page 写入 pages/ 目录。"""
        for rel_path, content in pages.items():
            full_path = os.path.join(pages_dir, rel_path)
            os.makedirs(os.path.dirname(full_path), exist_ok=True)
            with open(full_path, 'w', encoding='utf-8') as f:
                f.write(content)

    def _clean_pages_dir(self, pages_dir: str) -> None:
        """清空 pages/ 目录中的旧文件。"""
        import shutil
        if os.path.isdir(pages_dir):
            shutil.rmtree(pages_dir)
        os.makedirs(pages_dir, exist_ok=True)

    def _print_summary(self, stats: BuildStats) -> None:
        """输出生成摘要信息。"""
        print(f"\n{'='*50}")
        print(f"生成完成!")
        print(f"  文件总数: {stats.total_files}")
        print(f"  目录总数: {stats.total_dirs}")
        print(f"  生成页面: {stats.total_pages}")
        print(f"  index.html: {stats.index_size_bytes:,} bytes")
        print(f"  pages/ 总大小: {stats.pages_size_bytes:,} bytes")
        if stats.file_type_counts:
            print(f"  文件类型分布:")
            for ext, count in sorted(stats.file_type_counts.items()):
                print(f"    {ext}: {count}")
        if stats.errors:
            print(f"  错误 ({len(stats.errors)}):")
            for err in stats.errors:
                print(f"    - {err}")
        print(f"{'='*50}")


# ============================================================
# 入口
# ============================================================

if __name__ == '__main__':
    builder = IndexBuilder()
    builder.build()
